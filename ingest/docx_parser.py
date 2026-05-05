import docx
from typing import Dict, Any

class DOCXParser:
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extracts text from a .docx file."""
        results = {
            "raw_text": "",
            "paragraphs": [],
            "tables": [],
            "error": None
        }

        try:
            doc = docx.Document(file_path)
            full_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
                    results["paragraphs"].append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                results["tables"].append(table_data)

            results["raw_text"] = "\n".join(full_text)
            
        except Exception as e:
            results["error"] = str(e)
            
        return results

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        parser = DOCXParser()
        print(parser.extract_text(sys.argv[1])["raw_text"][:500])
